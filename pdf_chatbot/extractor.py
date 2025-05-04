import os
import PyPDF2
from docx import Document as DocxDocument
from odf.opendocument import load
from odf import text, teletype
from striprtf.striprtf import rtf_to_text
import pandas as pd
from nltk.tokenize import word_tokenize
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import tkinter as tk
from tkinter import filedialog
from preprocessing import preprocess
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class pdf_Extractor:
    @staticmethod
    def read_file(file_path):
        """Read text from various file types."""
        if not os.path.exists(file_path):
            logger.error(f"File '{file_path}' does not exist.")
            return None

        file_ext = os.path.splitext(file_path)[1].lower()

        try:
            if file_ext == '.pdf':
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text = ''
                    for page in pdf_reader.pages:
                        text += page.extract_text() or ''
                    return text

            elif file_ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as file:
                    return file.read()

            elif file_ext == '.docx':
                doc = DocxDocument(file_path)
                text = ''
                for para in doc.paragraphs:
                    text += para.text + '\n'
                return text

            elif file_ext == '.odt':
                doc = load(file_path)
                text = ''
                for element in doc.getElementsByType(text.P):
                    text += teletype.extractText(element) + '\n'
                return text

            elif file_ext == '.rtf':
                with open(file_path, 'r', encoding='utf-8') as file:
                    rtf_content = file.read()
                    return rtf_to_text(rtf_content)

            else:
                logger.error(f"Unsupported file format: {file_ext}. Supported formats: .pdf, .txt, .docx, .odt, .rtf")
                return None

        except Exception as e:
            logger.error(f"Error reading file '{file_path}': {str(e)}")
            return None

    @staticmethod
    def is_duplicate(db, new_text, threshold=0.8):
        """Check if the new text is a duplicate by comparing it with all existing documents."""
        from app import Document as SQLDocument  # Lazy import
        try:
            existing_texts = [doc.cleaned_text for doc in db.session.query(SQLDocument).all()]
            if not existing_texts:
                return False
            texts_to_compare = existing_texts + [new_text]
            vectorizer = TfidfVectorizer(stop_words='english')
            tfidf_matrix = vectorizer.fit_transform(texts_to_compare)
            if tfidf_matrix.shape[0] < 2 or tfidf_matrix.shape[1] == 0:
                return False
            similarities = cosine_similarity(tfidf_matrix[-1], tfidf_matrix[:-1])
            for score in similarities[0]:
                if score >= threshold:
                    return True
            return False
        except Exception as e:
            logger.error(f"Error checking for duplicates: {str(e)}")
            raise

    @staticmethod
    def process_file(file_path, chat_id, db):
        """Process a file, check for duplicates, and save to the database if unique."""
        from app import Document as SQLDocument  # Lazy import
        try:
            raw_text = pdf_Extractor.read_file(file_path)
            if raw_text is None:
                return False, "Failed to read file (unsupported extension or file error)", None
            cleaned_text = preprocess.clean_text(raw_text)
            if pdf_Extractor.is_duplicate(db, cleaned_text):
                return False, "Document is too similar to an existing one", None
            new_doc = SQLDocument(chat_id=chat_id, cleaned_text=cleaned_text)
            db.session.add(new_doc)
            db.session.commit()
            logger.info(f"Document processed and saved for chat_id: {chat_id}")
            return True, "Success", raw_text
        except Exception as e:
            logger.error(f"Error in process_file: {str(e)}")
            raise

    @staticmethod
    def select_file():
        """Open a file dialog to select a document."""
        file_path = filedialog.askopenfilename(
            title="Select a document",
            filetypes=[("All supported files", "*.pdf *.txt *.docx *.odt *.rtf")]
        )
        return file_path

    @staticmethod
    def is_query_relevant(db, query, similarity_threshold=0.05):
        """Check if a query is relevant to uploaded documents."""
        from app import Document as SQLDocument  # Lazy import
        try:
            if not query:
                return False, "Please input something"
            existing_texts = [doc.cleaned_text for doc in db.session.query(SQLDocument).all()]
            if not existing_texts:
                return False, "No documents have been uploaded"
            query_tokens = word_tokenize(query.lower())
            query_word_count = len([t for t in query_tokens if t.isalnum()])
            if query_word_count <= 10:
                query_keywords = ['summarize', 'what', 'explain', 'describe', 'how', 'why', 'pdf', 'document', 'is', 'are', 'define', 'overview']
                vectorizer = TfidfVectorizer(stop_words='english', max_features=50)
                tfidf_matrix = vectorizer.fit_transform(existing_texts)
                doc_keywords = vectorizer.get_feature_names_out()
                keyword_reference = ' '.join(query_keywords + list(doc_keywords))
                vectorizer = TfidfVectorizer(stop_words='english')
                tfidf_vectors = vectorizer.fit_transform([keyword_reference, query])
                if tfidf_vectors.shape[1] == 0:
                    return False, "Keyword reference is too sparse to process"
                similarity = cosine_similarity(tfidf_vectors[1], tfidf_vectors[0])[0][0]
                if similarity >= similarity_threshold:
                    return True, f"Short query is relevant (similarity: {similarity:.2f})"
                return False, "Short query is unrelated to documents"
            vectorizer = TfidfVectorizer(stop_words='english')
            tfidf_corpus = vectorizer.fit_transform(existing_texts)
            tfidf_query = vectorizer.transform([query])
            similarities = cosine_similarity(tfidf_query, tfidf_corpus)
            max_similarity = similarities[0].max()
            if max_similarity >= similarity_threshold:
                return True, f"Query is relevant (similarity: {max_similarity:.2f})"
            else:
                return False, "Query is unrelated to documents"
        except Exception as e:
            logger.error(f"Error checking query relevance: {str(e)}")
            return False, f"Error processing query: {str(e)}"

    @staticmethod
    def tokenize_to_df(text):
        if not pdf_Extractor.uploaded_texts:
            return False
        """Tokenize text and return a DataFrame with tokens."""
        if not text:
            return pd.DataFrame(columns=['Token'])
        
        # Tokenize the text
        tokens =word_tokenize(text)
        
        # Create DataFrame with tokens
        df = pd.DataFrame(tokens, columns=['Token'])
        return df